import pandas as pd
import matplotlib.pyplot as plt
import os
import numpy as np
import docx
from docx.shared import Mm

#Set the working directory to the current directory(script location)
path = os.path.dirname(os.path.realpath(__file__))
os.chdir(path)

#System and game details to include to report
system ='''
System details:

Operating system: Microsoft Windows 11 Pro, Version 10.0.22631
DirectX runtime version: DirectX 12
Driver: Game Ready Driver - 566.14 - Tue Nov 12, 2024
CPU: 13th Gen Intel(R) Core(TM) i9-13980HX
RAM: 32.0 GB
Storage: SSD - 1.9 TB

Graphics card
GPU processor: NVIDIA GeForce RTX 4090 Laptop GPU
Graphics clock: 2040 MHz
Memory bandwidth: 576.064 GB/s
Dedicated video memory: 16376 MB GDDR6

Display (1): LG Electronics LG ULTRAGEAR
Resolution: 2560 x 1440
Refresh rate: 144 Hz
'''
title_tested = '''
Tested title:

World of Warships
Build ID: 16989406
Graphics quality: Maximum preset

'''

#File with metrics to parse
file = 'menu_gameplay.hml'
if not os.path.exists(file):
    raise FileNotFoundError(f"File '{file}' not found.")

#Time intervals for plots
sec_menu = 25
sec_start_gamepaly = 100
sec_end_gameplay = 300

#Columns from the parsed data to include to report
columns_to_plot = ['GPU1 temperature, °C', 'GPU1 usage, %', 'CPU temperature, °C', 'CPU usage, %','RAM usage, MB','Framerate, FPS','HDD1 usage, %','NET3 download rate, MB/s']

#Parse data from the file with metrics
#Skip the first 2 rows with the system information
values = pd.read_csv(file, encoding='iso-8859-1', header = 2)
#Parse the column with the units for values
units = values.iloc[:,3]
units = list(units)
#Clean whitespaces for units
units = [i.strip() for i in units]
#Get the list of columns names
columns = list(values)
#Rename the first 2 columns with index and timestamp
columns[0] = '-'
columns[1] = 'timestamp'
#Clean whitespaces for column names
columns_names = [i.strip() for i in columns]
#Add units to columns names, skip first 2 columns
columns_names = [(i + ', '+ units[num-2]) if num > 1 else i for num, i in enumerate(columns_names)]
#Parse the data with the new column names
data = pd.read_csv(file, encoding='iso-8859-1', skiprows=len(columns_names)+1, header = None, names = columns_names)
#Get the framerate data from the parsed data
frameratedata = data.loc[:, data.columns.str.contains('Framerate')]
#Get the number of data rows
xlen = len(data)

#Get the main test results to add to the report
MinFramerate = 'Min Framerate, FPS: '+ str(round(np.min(data.loc[:,'Framerate, FPS']),2))
MaxFramerate = 'Max Framerate, FPS: '+ str(round(np.max(data.loc[:,'Framerate, FPS']),2))
MinGPU = 'Min GPU1 usage, %: '+ str(round(np.min(data.loc[:,'GPU1 usage, %']),2))
MaxGPU = 'Max GPU1 usage, %: '+ str(round(np.max(data.loc[:,'GPU1 usage, %']),2))
MinCPU = 'Min CPU usage, %: '+ str(round(np.min(data.loc[:,'CPU usage, %']),2))
MaxCPU = 'Max CPU usage, %: '+ str(round(np.max(data.loc[:,'CPU usage, %']),2))

#Class for creating plots
class CreatePlot:
    
    def __init__(self, i, lim_max, lim_min):
        #Data from column
        self.i = i
        #End time(position of data entity in the column) interval for plot
        self.lim_max = lim_max
        #Start time(position of data entity in the column) interval for plot
        self.lim_min = lim_min
    
    #Method to create a plot for current metric
    def create(self):
        #Create x axis with values from 0 to number of data rows witth step 1
        x = np.arange(0,xlen,1)
        #Create y axis with values from current column
        y = data[self.i]
        #Create plot figure and names
        plt.figure()
        plt.title(i)
        plt.ylabel(i)
        plt.xlabel("Time, Sec")
        #Create plot
        plt.plot(x,y)
        #Limit plot with time interval(positions of data entity in the column)
        plt.xlim(self.lim_min, self.lim_max)
        #Add min, max and average values limited by time interval(positions of data entity in the column)
        limited_value = (list(y))[self.lim_min:self.lim_max]
        max_value = round(np.max(limited_value),2)
        min_value = round(np.min(limited_value),2)
        average_value = round(np.mean(limited_value),2)
        plt.figtext(0.2, 0.04, f"Max = " + str(max_value) + "  Min = " + str(min_value), wrap=True, horizontalalignment='center', fontsize=10)
        plt.figtext(0.2, 0.01, f"Average = " + str(average_value), wrap=True, horizontalalignment='center', fontsize=10)
        #Name plot
        plot_name = ('plots/'+i.replace("/", "-") + ' ' + str(self.lim_min) + ' - ' + str(self.lim_max)+'.png')
        #Save plot as file
        plt.savefig(plot_name)
        #Add plot to the report
        report.add_picture(plot_name, height=Mm(95))
    
    #Method to normalize data for comparison plot(for clear view)
    def norm(self, tempdata):
        return (tempdata)/(np.max(tempdata)-np.min(tempdata)) if ((np.max(tempdata)) != (np.min(tempdata))) else tempdata
    
    #Method to create a comparison plot for current metric and framerate
    def createfr(self):
        #Create x axis with values from 0 to number of data rows witth step 1    
        x = np.arange(0,xlen,1)
        #Create y axis with values from current column
        y = data[self.i]
        #Create z(for overlay on y) axis with values from framerate column
        z = frameratedata
        #Skip framerate column
        if y.name != (z.columns[0]):
            #Create plot figure and names
            plt.figure()
            plt.title(i + ' and Framerate, FPS')
            plt.xlabel("Time, Sec")
            #Hide y(z) axis units
            plt.yticks(visible=False)
            #Create plots for current data and framerate with normalized values
            plt.plot(x,self.norm(y), label = i)
            plt.plot(x,self.norm(z), label = 'Framerate, FPS')
            plt.xlim(self.lim_min, self.lim_max)
            #Add legends for current data and framerate plots
            plt.legend()
            #Name plot
            plot_name = ('plots/'+i.replace("/", "-") + ' and Framerate, FPS ' + str(self.lim_min) + ' - ' + str(self.lim_max)+'.png')
            #Save plot as file
            plt.savefig(plot_name)
            #Add plot to the report
            report.add_picture(plot_name, height=Mm(95))
        
#Create directory for plots in current directory(script location)
plot_directory = path+'/plots'
try:
    os.mkdir(plot_directory)
except FileExistsError:
    pass

#Create the report document
report = docx.Document()
#Add general information to the report
report.add_heading("Performance report", 0)
report.add_paragraph(system)
report.add_paragraph(title_tested)
#Add main test results to the report
par = report.add_paragraph('Main results:')
par.runs[0].bold = True
report.add_paragraph(f'''
{MinFramerate}
{MaxFramerate}
{MinGPU}
{MaxGPU}
{MinCPU}
{MaxCPU}''')

#Create plots for data from columns
for num, i in enumerate(data):
    #Skip first 2 columns with index and timestamp
    if num > 1 and i in columns_to_plot:
        #Start the new report page
        report.add_page_break()
        #Add heading with the name of the column
        report.add_heading(i, 1)
        #Create plots using CreatePlot Class and createfr, create methods for all data(from 0 to sec_end_gameplay)
        plot = CreatePlot(i, sec_end_gameplay, 0)
        report.add_paragraph('Menu and Gameplay')
        plot.createfr()
        plot.create()
        #Create plots for data captured in menu(from 0 to sec_menu)
        plot = CreatePlot(i, sec_menu, 0)
        report.add_page_break()
        report.add_paragraph('Menu')
        plot.create()
        #Create plots for data captured during gameplay(from sec_start_gamepaly to sec_end_gameplay)
        plot = CreatePlot(i, sec_end_gameplay, sec_start_gamepaly)
        report.add_paragraph('Gameplay')
        plot.create()

        
    else:
        continue

#Save the report
report.save("Performance report.docx")